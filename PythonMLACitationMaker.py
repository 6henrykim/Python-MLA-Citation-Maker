"""
-------------------------------------------------------------------------------------------------------------------------------------------
    Creates a Word Doc with citations from an Excel spreadsheet specified by the user
    
    Author. "Title of source." italics(Title of container), Other contributors,
        Version, Number, Publisher, Publication date, Location.

-------------------------------------------------------------------------------------------------------------------------------------------
"""

#import libraries
from docx import Document
from docx.shared import Pt
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Length
from openpyxl import Workbook
from openpyxl import load_workbook
import datetime

"""
-------------------------------------------------------------------------------------------------------------------------------------------
    Row and Column numbers in Excel
-------------------------------------------------------------------------------------------------------------------------------------------
"""
ROW_DATA_STARTS = 2
#rows with data starts at 2

COL_NUM_AUTHORS = 1
COL_AUTH1_LAST_NAME = 2
COL_AUTH1_FIRST_NAME = 3
COL_AUTH1_MIDDLE_NAME = 4
COL_AUTH2_LAST_NAME = 5
COL_AUTH2_FIRST_NAME = 6
COL_AUTH2_MIDDLE_NAME = 7
COL_TITLE = 8
COL_CONTAINER = 9
COL_CONTRIBUTORS = 10
COL_VERSION = 11
COL_NUMBER = 12
COL_PUBLISHER = 13
COL_DATE_PUBLISHED = 14
COL_LOCATION = 15
COL_DATE_ACCESSED = 16
#to access data:
#sheet.cell(row, COLUMN).value



"""
-------------------------------------------------------------------------------------------------------------------------------------------
    Class that holds the attributes for each citation and the functions for getting
    that data from spreadsheet
-------------------------------------------------------------------------------------------------------------------------------------------
"""
class Citation:
    def __init__(self, worksheet):
        self.sheet = worksheet       #the excel sheet  
        self.numAuthors = 0          #number of authors add et al. if more than 2
        self.authors = ""            #list of authors in string format
        self.title = ""              #title of article
        self.container = ""          #title of collection or website, in italics
        self.contributors = ""       #editors or translators, Edited by ...
        self.version = ""            #edition or version
        self.number = ""             #number or vol, could be episode number
        self.publisher = ""          #publisher
        self.datePublished = ""      #date published or updated 
        self.location = ""           #page numbers or url, no https://
        self.dateAccessed = ""       #date website accessed
        self.sortKey = ""            #Author, title, and container used to sort citations

    #Read from Excel and store authors 
    def readAuthors(self, row):
        
        #read the number of authors user wants to record
        self.numAuthors = self.sheet.cell(row, COL_NUM_AUTHORS).value

        #get data from each column
        author1LastName = self.sheet.cell(row, COL_AUTH1_LAST_NAME).value
        author1FullName = str(author1LastName).capitalize()
        #if the first name column is filled add it to full name
        author1FirstName = self.sheet.cell(row, COL_AUTH1_FIRST_NAME).value
        if author1FirstName != None:         
            author1FullName += ", " + str(author1FirstName).capitalize()
            #if the middle name column is filled add it to full name
            author1MiddleName = self.sheet.cell(row, COL_AUTH1_MIDDLE_NAME).value
            if author1MiddleName != None:
                author1FullName += " " + str(author1MiddleName).capitalize()

        #if the first name column is filled add it to full name
        author2FirstName = self.sheet.cell(row, COL_AUTH2_FIRST_NAME).value
        if author2FirstName != None:    
            author2FullName = str(author2FirstName).capitalize()
        #if the middle name column is filled add it to full name
        author2MiddleName = self.sheet.cell(row, COL_AUTH2_MIDDLE_NAME).value
        if author2MiddleName != None:
            author2FullName += " " + str(author2MiddleName).capitalize()
        #if the last name column is filled add it to full name
        author2LastName = self.sheet.cell(row, COL_AUTH2_LAST_NAME).value
        if author2LastName != None:
            author2FullName += " " + str(author2LastName).capitalize()

        #assemble final string according to number of authors
        if self.numAuthors == None or self.numAuthors == 0:
            self.authors = ""
        elif self.numAuthors == 1:
            #one author format: LastName1, FirstName1 MiddleName1.
            self.authors = author1FullName + "."
        elif self.numAuthors == 2:
            #two author format: LastName1, FirstName1 MiddleName1 and FirstName2 MiddleName2 LastName2.
            self.authors = author1FullName + " and " + author2FullName + "."
        else:
            #more than 2 authors format: LastName1, FirstName1 MiddleName1, et al.
            self.authors = author1FullName + ", et al."

    #Read from Excel and store title
    def readTitle(self, row):
        
        #read the data
        uncapitalizedTitle = self.sheet.cell(row, COL_TITLE).value

        #check that the title column wasn't empty
        if uncapitalizedTitle != None:
            #add the opening quote
            self.title = "\""
            #capitalize the title and store it in the variable
            self.title += capitalizeTitle(uncapitalizedTitle)
            #add the closing quote
            self.title += ".\""
           
    #Read from Excel and store container name
    def readContainer(self, row):
        
        #read the data
        uncapitalizedContainer = self.sheet.cell(row, COL_CONTAINER).value

        #check that the container column wasn't empty
        if uncapitalizedContainer != None:
            #set the container variable to the capitalized version
            self.container = capitalizeTitle(uncapitalizedContainer)

    #Read from Excel and store contributors
    def readContributors(self, row):
        #read the data
        contributorsCell = self.sheet.cell(row, COL_CONTRIBUTORS).value

        #check that the contributors column wasn't empty
        if contributorsCell != None:
            #set the contributors variable to the cell value
            self.contributors = contributorsCell

    #Read from Excel and store version
    def readVersion(self, row):
        #read the data
        versionCell = self.sheet.cell(row, COL_VERSION).value

        #check that the version column wasn't empty
        if versionCell != None:
            #set the version variable to the cell value
            self.version = versionCell
    
    #Read from Excel and store number
    def readNumber(self, row):
        #read the data
        numberCell = self.sheet.cell(row, COL_NUMBER).value

        #check that the number column wasn't empty
        if numberCell != None:
            #set the contributors variable to the cell value
            self.number = str(numberCell)
            
    #Read from Excel and store publisher
    def readPublisher(self, row):
        #read the data
        publisherCell = self.sheet.cell(row, COL_PUBLISHER).value

        #check that the publisher column wasn't empty
        if publisherCell != None:
            #set the publisher variable to the cell value
            self.publisher = capitalizeTitle(publisherCell)
                  
    
    #Read from Excel and store published date
    def readDatePublished(self, row):
        datePublishedCell = self.sheet.cell(row, COL_DATE_PUBLISHED).value
        #if date is in datetime format convert into a string
        if(type(datePublishedCell) == datetime.datetime):
            self.datePublished = str(datePublishedCell.day) + " " + convertNumToMonth(datePublishedCell.month) + " " + str(datePublishedCell.year)
        #make sure date column isn't empty
        elif datePublishedCell == None:
            self.datePublished = ""
        #otherwise store it as a string
        else:
            self.datePublished = str(datePublishedCell)


    #Read from Excel and store location
    def readLocation(self, row):
        #read the data
        locationCell = self.sheet.cell(row, COL_LOCATION).value

        #check that the version column wasn't empty
        if locationCell != None:
            #if url has http:// or https://, only store the part after it
            if locationCell.startswith("http://") or locationCell.startswith("https://"):
                urlSplit = locationCell.split("//")
                self.location = urlSplit[1]
            #otherwise store the text as is
            else:
                self.location = locationCell

            
    #Read from Excel and store accessed date
    def readDateAccessed(self, row):
        dateAccessedCell = self.sheet.cell(row, COL_DATE_ACCESSED).value
        #if date is in datetime format convert into a string
        if(type(dateAccessedCell) == datetime.datetime):
            self.dateAccessed = "Accessed " + str(dateAccessedCell.day) + " " + convertNumToMonth(dateAccessedCell.month) + " " + str(dateAccessedCell.year)
        #make sure date column isn't empty
        elif dateAccessedCell == None:
            self.dateAccessed = ""
        #otherwise store it as a string
        else:
            self.dateAcessed = "Accessed " + str(dateAccessedCell)
    

    #Create the sorting key based on data read
    def createSortKey(self):
        self.sortKey += self.authors + " " + self.title + " " + self.container
        self.sortKey = self.sortKey.lower()
        

"""
-------------------------------------------------------------------------------------------------------------------------------------------
    Function to get the name of the output file from the user
-------------------------------------------------------------------------------------------------------------------------------------------
"""
def inputExcelFileName():
    
    fileExtension = ".xlsx"
    
    #get a name of file from user
    fileName = input("Enter name of Excel file or its file path: ")

    #remove any quote marks around the name
    if fileName.startswith("\""):
        fileNameSplit = fileName.split("\"")
        fileName = fileNameSplit[1]
    if fileName.endswith("\""):
        fileNameSplit = fileName.split("\"")
        fileName = fileNameSplit[0]
        
    #append .xlsx if input name doesn't have it
    if (fileName.endswith(fileExtension) == False):
        fileName += fileExtension

    #return the name of the file
    return fileName



"""
-------------------------------------------------------------------------------------------------------------------------------------------
    Function to get the name of the output file from the user
-------------------------------------------------------------------------------------------------------------------------------------------
"""
def inputCitationFileName():
    
    #default name of the document
    fileName = "citations.docx"
    fileExtension = ".docx"

    #get a name of document from user
    fileName = input("Enter name of Word Doc or file path to save to: ")

    #remove any quote marks around the name
    if fileName.startswith("\""):
        fileNameSplit = fileName.split("\"")
        fileName = fileNameSplit[1]
    if fileName.endswith("\""):
        fileNameSplit = fileName.split("\"")
        fileName = fileNameSplit[0]
    
    #append .docx if input name doesn't have it
    if (fileName.endswith(fileExtension) == False):
        fileName += fileExtension

    #return the name of the file
    return fileName



"""
-------------------------------------------------------------------------------------------------------------------------------------------
    Function to return a month string based on number
-------------------------------------------------------------------------------------------------------------------------------------------
"""
def convertNumToMonth(num):
    if num == 1:
        return "Jan."
    elif num == 2:
        return "Feb."
    elif num == 3:
        return "Mar."
    elif num == 4:
        return "Apr."
    elif num == 5:
        return "May"
    elif num == 6:
        return "Jun."
    elif num == 7:
        return "Jul."
    elif num == 8:
        return "Aug."
    elif num == 9:
        return "Sep."
    elif num == 10:
        return "Oct."
    elif num == 11:
        return "Nov."
    elif num == 12:
        return "Dec."
    else:
        return "???"

"""
-------------------------------------------------------------------------------------------------------------------------------------------
    Function to return a capitalized string according to title conventions
-------------------------------------------------------------------------------------------------------------------------------------------
"""
def capitalizeTitle(uncapitalizedString):

    capitalizedTitle = ""
    
    #list of words that shouldn't be capitalized including articles, prepositions, and coordinate conjunctives
    dontCapitalize = ["the", "a", "an", "with", "for", "and", "nor", "but", "or", "yet", "so", "at", "around", "by", "after", "along", "for", "from", "of", "on", "to", "with", "without"]

    #set up list to hold individual words
    individualWords = uncapitalizedString.split(" ")
    
    #capitalize words, loop for each word
    for i in range(0, len(individualWords)):
        #always capitalize first word
        if i == 0:
            individualWords[i] = individualWords[i].capitalize()
        else:
            #set flag for whether the word should be capitalize
            wordShouldCapitalize = True
            
            #loop for list of words that shouldn't be capitalized and check if the word matches any of them
            for j in range(0, len(dontCapitalize)):
                if individualWords[i] == dontCapitalize[j]:
                    wordShouldCapitalize = False
                
            if wordShouldCapitalize:
                individualWords[i] = individualWords[i].capitalize()

        
        #add the words to the title
        capitalizedTitle += individualWords[i]
        #add spaces in between the words
        if i < len(individualWords) - 1:
            capitalizedTitle += " "

    return capitalizedTitle


"""
-------------------------------------------------------------------------------------------------------------------------------------------
    Function to format citation paragraphs to be double space and hanging indent
-------------------------------------------------------------------------------------------------------------------------------------------
"""
def formatParagraph(paragraph):
    paragraphFormatting = paragraph.paragraph_format
    paragraphFormatting.line_spacing_rule = 2   #set double spaceing
    paragraphFormatting.keep_together = True    #keeps citation on same page
    paragraphFormatting.left_indent = Inches(0.5)   #indent citations
    paragraphFormatting.first_line_indent = Inches(-0.5)    #negative to make fist line hanging indent
    


"""
-------------------------------------------------------------------------------------------------------------------------------------------
    Function to format heading and citation runs to be Times New Roman and point 12
-------------------------------------------------------------------------------------------------------------------------------------------
"""
def formatRun(run):
    font = run.font
    font.name = "Times New Roman"
    font.size = Pt(12)

    
"""
-------------------------------------------------------------------------------------------------------------------------------------------
    Function to determine what should go before the next string in an MLA citation
-------------------------------------------------------------------------------------------------------------------------------------------
"""
def determineTransition(stringToCheck, needComma):
    #if stringToCheck will be the first thing written, nothing goes before it
    if stringToCheck == "":
        return ""
    #elif stringToCheck doesnt end with the authors or title, a comma goes before it 
    elif needComma:
        return ", "
    else:
        return " "


"""
-------------------------------------------------------------------------------------------------------------------------------------------
    Function to sort the list of citations by alphabetical order using bubble sort
-------------------------------------------------------------------------------------------------------------------------------------------
"""
def sortCitations(citationList):
    print("todo")
    
    maxElement = len(citationList) - 1
    while maxElement > 0:
        
        index = 0
        while index < maxElement:

            #check if the first element comes after the second one and swap if it does
            if citationList[index].sortKey > citationList[index + 1].sortKey:
                #swap the elements
                temp = citationList[index]
                citationList[index] = citationList[index + 1]
                citationList[index + 1] = temp


            index += 1


            
        maxElement -=1
    

    
"""
-------------------------------------------------------------------------------------------------------------------------------------------
    Main function execution
-------------------------------------------------------------------------------------------------------------------------------------------
"""

print("MLA Citation Maker")    

#flags for whether the files opened properly
excelFileOpened = True
citationFileOpened = True

excelFileName = inputExcelFileName()
#try to open the Excel file
try:
    #load workbook 
    workbook = load_workbook(filename = excelFileName)
    #get the name of the first worksheet
    sheetName = workbook.sheetnames[0]
    #set the sheet using the name of the first worksheet
    worksheet = workbook[sheetName]
except:
    excelFileOpened = False
    print("Failed to open \"" + excelFileName + "\"")
    #pause for user to see error message
    input("Press Enter to quit")


#Try to open the Word doc if the Excel file opened
if excelFileOpened:
    
    documentName = inputCitationFileName()
    #try to open the specified file and save it to make sure it's not being used by another program
    try:
        document = Document(documentName)
        document.save(documentName)
    #display error message if file couldn't be opened and saved
    except PermissionError:
        citationFileOpened = False
        print("Failed to open \"" + documentName + "\": make sure no program has the file open")
        #pause for user to see error message
        input("Press Enter to quit")
    #otherwise create the file
    except:
        document = Document()
        document.save(documentName)

#Continue if both files opened
if excelFileOpened and citationFileOpened:
    print("Succesfully accessed both files")

    #Write the heading to the citation doc
    heading = document.add_paragraph()
    headerFormatting = heading.paragraph_format
    headerFormatting.line_spacing_rule = 2                      #set double spaceing
    headerFormatting.alignment = WD_ALIGN_PARAGRAPH.CENTER      #center the heading
    headerFormatting.page_break_before = True                   #put header on a new page
    headingRun = heading.add_run("Works Cited")
    formatRun(headingRun)

    #list to hold each citation object
    citationList = []
    
    #loop to read from sheet as long as the first author's last name, title, or container is filled
    index = ROW_DATA_STARTS 
    while (worksheet.cell(index, COL_AUTH1_LAST_NAME).value != None) or (worksheet.cell(index, COL_TITLE).value != None) or (worksheet.cell(index, COL_CONTAINER).value != None):
        citation = Citation(worksheet)
        citation.readAuthors(index)
        citation.readTitle(index)
        citation.readContainer(index)
        citation.readContributors(index)
        citation.readVersion(index)
        citation.readNumber(index)
        citation.readPublisher(index)
        citation.readDatePublished(index)
        citation.readLocation(index)
        citation.readDateAccessed(index)
        citation.createSortKey()
        citationList.append(citation)
        index += 1

    #Sort the citations by alphabetical order
    sortCitations(citationList)
    
    #Loop through each citation and write to document
    for citation in citationList:
        #tracks everything that was written to document
        stringWritten = ""
        #flag for whether the next segment should have a comma before it
        needComma = False
        #add the paragraph for each citation
        citationParagraph = document.add_paragraph()
        formatParagraph(citationParagraph)

        #Check if there's content for each segment and add it as a run if there is
        if citation.authors != "":
            citationRun = citationParagraph.add_run(citation.authors)
            formatRun(citationRun)
            stringWritten += citation.authors
            
        if citation.title != "":
            transition = determineTransition(stringWritten, needComma)
            citationRun = citationParagraph.add_run(transition + citation.title)
            formatRun(citationRun)
            stringWritten += transition + citation.title

        if citation.container != "":
            transition = determineTransition(stringWritten, needComma)
            citationRun = citationParagraph.add_run(transition + citation.container)
            formatRun(citationRun)
            citationRun.font.italic = True
            stringWritten += transition + citation.container
            needComma = True

        if citation.contributors != "":
            transition = determineTransition(stringWritten,needComma)
            citationRun = citationParagraph.add_run(transition + citation.contributors)
            formatRun(citationRun)
            stringWritten += transition + citation.contributors
            needComma = True
            
        if citation.version != "":
            transition = determineTransition(stringWritten, needComma)
            citationRun = citationParagraph.add_run(transition + citation.version)
            formatRun(citationRun)
            stringWritten += transition + citation.version
            needComma = True

        if citation.number != "":
            transition = determineTransition(stringWritten, needComma)
            citationRun = citationParagraph.add_run(transition + citation.number)
            formatRun(citationRun)
            stringWritten += transition + citation.number
            needComma = True

        if citation.publisher != "":
            transition = determineTransition(stringWritten, needComma)
            citationRun = citationParagraph.add_run(transition + citation.publisher)
            formatRun(citationRun)
            stringWritten += transition + citation.publisher
            needComma = True
    

        if citation.datePublished != "":
            transition = determineTransition(stringWritten, needComma)
            citationRun = citationParagraph.add_run(transition + citation.datePublished)
            formatRun(citationRun)
            stringWritten += transition + citation.datePublished
            needComma = True
        
        if citation.location != "":
            transition = determineTransition(stringWritten, needComma)
            citationRun = citationParagraph.add_run(transition + citation.location)
            formatRun(citationRun)
            stringWritten += transition + citation.location
            needComma = True
            
        if citation.dateAccessed != "":
            transition = determineTransition(stringWritten, needComma)
            citationRun = citationParagraph.add_run(transition + citation.dateAccessed)
            formatRun(citationRun)
            stringWritten += transition + citation.dateAccessed
            needComma = True

        #Add a period to the end if there isn't one already
        if (stringWritten.endswith(".") == False):
            citationRun = citationParagraph.add_run(".")
            formatRun(citationRun)
            stringWritten += "."
            

        print(stringWritten)
        
    print("Citation page complete. Warning: Acryonyms may not be capitalized correctly")
    input("Press Enter to quit")



    #Save the file
    document.save(documentName)


    




